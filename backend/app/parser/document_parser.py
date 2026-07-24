import os
import re
import io
import requests
from bs4 import BeautifulSoup
from app.parser.code_parser import MarkdownParser

class PDFParser:
    """
    Parses PDF documents (.pdf) into structured Markdown sections for RAG indexing.
    """
    @staticmethod
    def parse(source: str | bytes, rel_path: str = "document.pdf") -> list[dict]:
        try:
            import pypdf
        except ImportError:
            print("[PDFParser] pypdf library is not installed.")
            return []

        try:
            if isinstance(source, bytes):
                reader = pypdf.PdfReader(io.BytesIO(source))
            else:
                if not os.path.exists(source):
                    return []
                reader = pypdf.PdfReader(source)

            extracted_text = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    extracted_text.append(f"## Page {i + 1}\n\n{page_text.strip()}")

            full_md = "\n\n".join(extracted_text)
            if not full_md.strip():
                return []

            chunks = MarkdownParser.parse(full_md, rel_path)
            for c in chunks:
                c['chunk_type'] = 'documentation'
            return chunks
        except Exception as e:
            print(f"[PDFParser] Error parsing PDF {rel_path}: {e}")
            return []


class DocxParser:
    """
    Parses Microsoft Word documents (.docx) into structured Markdown sections.
    """
    @staticmethod
    def parse(source: str | bytes, rel_path: str = "document.docx") -> list[dict]:
        try:
            import docx
        except ImportError:
            print("[DocxParser] python-docx library is not installed.")
            return []

        try:
            if isinstance(source, bytes):
                doc = docx.Document(io.BytesIO(source))
            else:
                if not os.path.exists(source):
                    return []
                doc = docx.Document(source)

            extracted_lines = []
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                style = p.style.name.lower() if p.style else ""
                if "heading 1" in style:
                    extracted_lines.append(f"# {text}")
                elif "heading 2" in style:
                    extracted_lines.append(f"## {text}")
                elif "heading 3" in style:
                    extracted_lines.append(f"### {text}")
                else:
                    extracted_lines.append(text)

            # Process tables in docx
            for table in doc.tables:
                table_md = []
                for r_idx, row in enumerate(table.rows):
                    cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                    table_md.append("| " + " | ".join(cells) + " |")
                    if r_idx == 0:
                        table_md.append("| " + " | ".join(["---"] * len(cells)) + " |")
                if table_md:
                    extracted_lines.append("\n" + "\n".join(table_md) + "\n")

            full_md = "\n\n".join(extracted_lines)
            if not full_md.strip():
                return []

            chunks = MarkdownParser.parse(full_md, rel_path)
            for c in chunks:
                c['chunk_type'] = 'documentation'
            return chunks
        except Exception as e:
            print(f"[DocxParser] Error parsing DOCX {rel_path}: {e}")
            return []


class WebPageParser:
    """
    Scrapes public web pages and Google Docs live links, returning structured Markdown documentation chunks.
    """
    @staticmethod
    def scrape(url: str, custom_title: str = "") -> list[dict]:
        if not url or not (url.startswith("http://") or url.startswith("https://")):
            return []

        headers = {
            "User-Agent": "ZeroTicket-DocBot/1.0 (+https://github.com/jiayong1008/zero-ticket)"
        }

        # 1. Detect Google Docs URLs and convert to live text export
        gdoc_match = re.search(r'docs\.google\.com/document/d/([a-zA-Z0-9_-]+)', url)
        if gdoc_match:
            doc_id = gdoc_match.group(1)
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
            try:
                resp = requests.get(export_url, headers=headers, timeout=10)
                if resp.status_code == 200 and resp.text.strip():
                    doc_title = custom_title or f"Google Doc ({doc_id[:8]})"
                    content = f"# {doc_title}\n\n" + resp.text.strip()
                    rel_path = f"gdoc_{doc_id[:8]}.md"
                    chunks = MarkdownParser.parse(content, rel_path)
                    for c in chunks:
                        c['chunk_type'] = 'documentation'
                    return chunks
            except Exception as e:
                print(f"[WebPageParser] Google Doc export fetch failed: {e}")

        # 2. General Documentation Webpage Scraping
        try:
            resp = requests.get(url, headers=headers, timeout=10)
            if resp.status_code != 200:
                print(f"[WebPageParser] Webpage request returned status {resp.status_code}")
                return []

            soup = BeautifulSoup(resp.text, 'html.parser')

            # Strip non-content script, style, nav, header, footer elements
            for element in soup(["script", "style", "nav", "header", "footer", "aside", "svg"]):
                element.extract()

            # Extract main article title
            page_title = custom_title
            if not page_title:
                title_tag = soup.find('title')
                if title_tag and title_tag.string:
                    page_title = title_tag.string.strip()
                else:
                    page_title = url

            # Find main article body container if present
            main_container = soup.find('main') or soup.find('article') or soup.find(id=re.compile(r'content|body|main', re.I)) or soup.body
            if not main_container:
                main_container = soup

            lines = []
            for elem in main_container.find_all(['h1', 'h2', 'h3', 'p', 'li', 'pre', 'code']):
                text = elem.get_text().strip()
                if not text:
                    continue
                tag = elem.name.lower()
                if tag == 'h1':
                    lines.append(f"\n# {text}\n")
                elif tag == 'h2':
                    lines.append(f"\n## {text}\n")
                elif tag == 'h3':
                    lines.append(f"\n### {text}\n")
                elif tag == 'li':
                    lines.append(f"- {text}")
                elif tag in ['pre', 'code']:
                    lines.append(f"```\n{text}\n```")
                else:
                    lines.append(text)

            clean_md = f"# {page_title}\n\nSource URL: {url}\n\n" + "\n\n".join(lines)
            safe_domain = re.sub(r'[^a-zA-Z0-9]', '_', url.replace("https://", "").replace("http://", ""))[:30]
            rel_path = f"web_{safe_domain}.md"

            chunks = MarkdownParser.parse(clean_md, rel_path)
            for c in chunks:
                c['chunk_type'] = 'documentation'
            return chunks
        except Exception as e:
            print(f"[WebPageParser] Error scraping webpage {url}: {e}")
            return []
