import unittest
from app.parser.document_parser import PDFParser, DocxParser, WebPageParser

class TestDocumentParser(unittest.TestCase):
    def test_webpage_parser_google_doc_conversion(self):
        # Test scraping a public documentation page
        url = "https://example.com"
        chunks = WebPageParser.scrape(url, custom_title="Sample Documentation")
        self.assertGreater(len(chunks), 0)
        self.assertEqual(chunks[0]['chunk_type'], 'documentation')
        self.assertIn("file_path", chunks[0])

    def test_docx_parser(self):
        try:
            import docx
            doc = docx.Document()
            doc.add_heading("ZeroTicket Setup Manual", level=1)
            doc.add_paragraph("This is an automated setup test paragraph.")
            import io
            stream = io.BytesIO()
            doc.save(stream)
            bytes_data = stream.getvalue()

            chunks = DocxParser.parse(bytes_data, "test_manual.docx")
            self.assertGreater(len(chunks), 0)
            self.assertIn("ZeroTicket Setup Manual", chunks[0]['content'])
        except Exception as e:
            self.fail(f"DocxParser test failed with error: {e}")

if __name__ == '__main__':
    unittest.main()
